"""
Generates SF_Tech_Debt_Assessor_All_Checks_2026-06-17.docx
A complete reference of every check across all 22 categories.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Severity colours ────────────────────────────────────────────────
SEV_COLORS = {
    'critical': RGBColor(0xC0, 0x39, 0x2B),
    'high':     RGBColor(0xD3, 0x54, 0x00),
    'medium':   RGBColor(0xB8, 0x86, 0x0B),
    'low':      RGBColor(0x27, 0xAE, 0x60),
}
SEV_BG = {
    'critical': '5D1010',
    'high':     '6B2A00',
    'medium':   '5C4200',
    'low':      '0E4D28',
}
SEV_LABEL = {
    'critical': 'CRITICAL',
    'high':     'HIGH',
    'medium':   'MEDIUM',
    'low':      'LOW',
}

# ── All checks ──────────────────────────────────────────────────────
CATEGORIES = [
  {
    'name': 'Configuration',
    'checks': 13,
    'group': 'Configuration & Architecture',
    'checks_list': [
      ('Active Workflow Rules', 'high'),
      ('Active Process Builders', 'high'),
      ('Active Automation Components — Review for Overlap (>50 total)', 'medium'),
      ('Active Validation Rules (>50)', 'medium'),
      ('Validation Rules Without Descriptions', 'low'),
      ('Classic Approval Processes Still Active', 'medium'),
      ('Flows Using Legacy Einstein for Flow Actions', 'low'),
      ('Web-to-Case Enabled Without CAPTCHA — Spam Risk', 'high'),
      ('Active Case Auto-Response Rules — Legacy', 'low'),
      ('s-Controls Still Active — Deprecated Technology', 'critical'),
      ("Active PushTopics — Deprecated Summer '26", 'high'),
      ('Pending Time-Based Workflow Actions in Queue', 'medium'),
      ('No Login Flows Configured — Additional Auth Enforcement Missing', 'low'),
    ],
  },
  {
    'name': 'Code Quality',
    'checks': 29,
    'group': 'Code & Development',
    'checks_list': [
      ('Triggers with Business Logic (No Handler Pattern)', 'high'),
      ('Classes/Triggers Below 75% Test Coverage', 'high'),
      ('Components on Outdated API Versions (< v55)', 'medium'),
      ('SOQL Queries Inside Loops — Governor Limit Risk', 'critical'),
      ('Hardcoded Salesforce IDs in Classes', 'high'),
      ('Empty catch Blocks — Exceptions Silently Swallowed', 'high'),
      ('DML Operations Inside Loops — Governor Limit Risk', 'critical'),
      ('Schema.getGlobalDescribe() — Expensive Mass Schema Lookup', 'medium'),
      ('Apex Classes Without a Sharing Declaration', 'high'),
      ('System.setPassword() — AppExchange Security Violation', 'critical'),
      ('UserInfo.getSessionId() — Session ID Exposure Risk', 'high'),
      ('SOQL Queries Without FLS Enforcement (WITH SECURITY_ENFORCED / USER_MODE)', 'medium'),
      ("SOAP login() Usage — Retired Spring '26 Default, Hard Retirement Summer '27", 'medium'),
      ("Hardcoded login.salesforce.com URLs — My Domain Enforced Spring '26", 'high'),
      ('Dynamic SOQL with String Concatenation — SOQL Injection Risk', 'critical'),
      ('@future(callout=true) Methods That Also Perform DML', 'medium'),
      ('Weak Cryptographic Algorithms — MD5 / SHA-1 Usage', 'high'),
      ('@IsTest(SeeAllData=true) — Tests Access Real Org Data', 'high'),
      ('Test Classes With No Assert Statements', 'high'),
      ('Test Classes Missing Test.startTest() / Test.stopTest()', 'medium'),
      ('Test Classes Inserting Data Without @TestSetup', 'low'),
      ('Apex Classes Use the global Access Modifier (PMD: AvoidGlobalModifier)', 'high'),
      ('Queueable Classes Without a Finalizer (PMD: QueueableWithoutFinalizer)', 'medium'),
      ('Apex Classes Use the @future Annotation (PMD: AvoidFutureAnnotation)', 'medium'),
      ('DML Operations in Apex Constructors — CSRF-Class Vulnerability (PMD: ApexCSRF)', 'high'),
      ('addError() Called with escape=false — XSS Risk (PMD: ApexXSSFromEscapeFalse)', 'high'),
      ('HTTP (Non-HTTPS) Callout Endpoints (PMD: ApexInsecureEndpoint)', 'high'),
      ('System.debug Statements in Production Code (PMD: AvoidDebugStatements)', 'medium'),
      ('SOQL Queries Without WHERE or LIMIT — Full Table Scan Risk (PMD: AvoidNonRestrictiveQueries)', 'high'),
    ],
  },
  {
    'name': 'Data Model',
    'checks': 4,
    'group': 'Configuration & Architecture',
    'checks_list': [
      ('Custom Objects Without Descriptions', 'low'),
      ('Custom Fields Lack Descriptions (>50%)', 'medium'),
      ('Objects with 100+ Custom Fields', 'high'),
      ('Custom Object Count in Org (>200)', 'medium'),
    ],
  },
  {
    'name': 'Service Cloud',
    'checks': 69,
    'group': 'CRM & Service',
    'checks_list': [
      ('Excessive Case Record Types (>10)', 'medium'),
      ('Inactive Case Record Types', 'low'),
      ('Excessive Queues Configured (>50)', 'medium'),
      ('Case Assignment Rules — Legacy', 'medium'),
      ('Escalation Rules — Legacy', 'low'),
      ("Unverified Organization-Wide Email Addresses — Fail to Send Spring '26", 'high'),
      ('No Omni-Channel Service Channels Configured', 'medium'),
      ('Routing Configurations Use Legacy Tab-Based Capacity', 'critical'),
      ('Routing Configurations Have No Push Timeout', 'high'),
      ('No Skills-Based Routing — All Queues Use Availability-Only Routing', 'medium'),
      ('Presence Configurations Have No Capacity Limit', 'high'),
      ('No Presence Configurations Defined Despite Active Service Channels', 'medium'),
      ('Knowledge Enabled but No Published Articles Found', 'medium'),
      ('Draft Knowledge Articles Stalled for 180+ Days', 'medium'),
      ('Published Knowledge Articles Not Updated in 12+ Months', 'high'),
      ('No Data Category Groups Configured — Knowledge Cannot Be Filtered by Audience', 'high'),
      ('Published Articles Have No Data Category Assignment', 'high'),
      ('Published Articles Have No Validation Status', 'medium'),
      ('Active Entitlement Processes Have No Business Hours Assigned', 'critical'),
      ('Active Entitlement Processes Have No Milestone Actions', 'critical'),
      ('Open Cases With Entitlement but No SLA Start Date', 'high'),
      ('Service Contracts Have No Linked Entitlements', 'high'),
      ('Email-to-Case Routing Addresses Without TLS', 'critical'),
      ('Email-to-Case Routing Addresses Have No Default Owner', 'high'),
      ('Inbound Emails Creating New Cases Instead of Threading (Last 30 Days)', 'high'),
      ('Email Service Addresses Accept Emails from Any Sender', 'medium'),
      ('Live Chat Buttons Not Routed Through Omni-Channel', 'critical'),
      ('Active Legacy Live Agent Deployments Without Embedded Service Config', 'critical'),
      ('Legacy Live Chat Active but Messaging for In-App and Web (MIAW) Not Adopted', 'high'),
      ('Live Chat Buttons Point to Empty or Deleted Queues', 'high'),
      ('No Lightning Service Console App Detected', 'high'),
      ('No Active Macros Configured', 'medium'),
      ('No Einstein Next Best Action Recommendation Strategies Configured', 'low'),
      ('Call Center Configured but No Default Softphone Layout Assigned', 'medium'),
      ('Messaging Channels Without OPTOUT Keyword — TCPA/GDPR Risk', 'high'),
      ('Open Cases With Active SLA Milestone Violations', 'critical'),
      ('Open Escalated Cases With No Activity in 3+ Days', 'critical'),
      ('Open Cases Not Modified in 7+ Days', 'high'),
      ('High Zero-Touch Case Close Rate — No Documented Activity (Last 90 Days)', 'high'),
      ('Entitlements Past End Date Still Marked Active', 'high'),
      ('Open Cases Linked to Expired Entitlements', 'critical'),
      ('No Active Quick Texts Configured', 'medium'),
      ('Quick Texts Not Updated in 12+ Months (>50%)', 'low'),
      ('Callback Requests Unresolved After 24+ Hours', 'high'),
      ('Messaging Sessions Ended With No Agent Response (Last 30 Days)', 'high'),
      ('Completed Chat Transcripts With No Linked Case (Last 30 Days)', 'medium'),
      ('Inbound Social Posts With No Linked Case (Last 30 Days)', 'high'),
      ('No Case Team Templates Configured', 'low'),
      ('Active Entitlements Never Linked to Any Case', 'medium'),
      ('Cases Linked to Multiple Entitlements', 'medium'),
      ('Business Hours Records Exclude Weekends but Have No Holiday Exceptions', 'medium'),
      ('Entitlement Milestones With Target Time of 5 Minutes or Less', 'high'),
      ('Entitlement Processes Have Milestones With Duplicate Target Times', 'high'),
      ('Published Articles Still Visible in Legacy Customer Self-Service Portal Channel', 'low'),
      ('Knowledge Search Not Optimised — No Promoted Terms or Synonym Groups Configured', 'medium'),
      ('Duplicate Published Article Titles Found', 'medium'),
      ('Published Articles Have No Summary', 'low'),
      ('No Knowledge Articles Have Been Attached to Any Case (Zero Deflection Evidence)', 'medium'),
      ('Open Cases With No Contact and No Account', 'high'),
      ('Open Cases With No Priority Set', 'medium'),
      ('Open Cases With No Origin (Unknown Channel)', 'medium'),
      ('Open Cases Older Than 90 Days', 'high'),
      ('Open Cases With No Description', 'low'),
      ('Open Cases Owned by Individual Users (Not Queues)', 'medium'),
      ('Open Incidents With No Related Items', 'medium'),
      ('Open Swarms With No Activity in 14+ Days', 'low'),
      ('Open Work Orders With No Case and No Asset', 'medium'),
      ('Active CSAT Surveys Configured but No Survey Responses Received', 'medium'),
      ('Completed Voice Calls With No Linked Case (Last 30 Days)', 'high'),
    ],
  },
  {
    'name': 'Sharing & Security',
    'checks': 24,
    'group': 'Security & Access',
    'checks_list': [
      ('Objects with Public Read/Write OWD (admin-configurable objects only)', 'critical'),
      ('Objects with Public Read Only OWD (>5 admin-configurable objects)', 'medium'),
      ('Too Many Profiles Configured (>20)', 'medium'),
      ('Permission Sets Without Descriptions', 'low'),
      ('Excessive Sharing Rules Configured (>50)', 'high'),
      ('Active Users with Modify All Data', 'critical'),
      ('Active Users Inactive for 90+ Days', 'high'),
      ('Integration/API Users Without IP Restrictions', 'high'),
      ('Standard Profiles Without Login IP Restrictions', 'medium'),
      ('Active Users Not Enrolled in MFA (TwoFactorInfo) — Sandbox: Low Severity', 'critical'),
      ('Security Health Check Score Below Threshold (<75)', 'critical'),
      ('Active OAuth Access Tokens (>100)', 'medium'),
      ('Users with Standard-Assurance Sessions (No MFA Step-Up)', 'medium'),
      ('Users with Passwords That Never Expire', 'high'),
      ('Active Sites with Guest User Access', 'high'),
      ('Permission Sets Grant Privileged Access — Phishing-Resistant MFA Required (May 2026)', 'high'),
      ('Async Sharing Recalculation Release Update Not Activated', 'medium'),
      ('Guest Profiles with Case Read Access — Unauthenticated Data Exposure', 'critical'),
      ('Active Outbound Messages — Session IDs Retired February 2026', 'high'),
      ('Permission Set Groups Not in Use', 'medium'),
      ('Users Assigned 10+ Custom Permission Sets', 'medium'),
      ('Cloned System Administrator Profiles Detected', 'critical'),
      ('No Transaction Security Policies Configured', 'low'),
      ('Active Users (stale account risk proxy)', 'medium'),
    ],
  },
  {
    'name': 'Integrations',
    'checks': 9,
    'group': 'Configuration & Architecture',
    'checks_list': [
      ('Remote Sites with Protocol Security Disabled', 'critical'),
      ('Inactive Remote Site Settings', 'low'),
      ('Connected Apps Without Descriptions', 'low'),
      ('Apex Classes with Hardcoded HTTP Endpoints (Not Using Named Credentials)', 'high'),
      ('Named Credentials Using Per-User Authentication', 'medium'),
      ('Apex Classes on Retired API Versions (≤v30) — Broken Since Summer \'25', 'critical'),
      ("Active PushTopics — Streaming API Deprecated Summer '26", 'high'),
      ('No External Credentials Configured — Using Legacy Named Credentials Only', 'medium'),
      ('No Dedicated Integration User Profiles Found — Connected Apps May Use Named Users', 'medium'),
    ],
  },
  {
    'name': 'Test Coverage',
    'checks': 4,
    'group': 'Code & Development',
    'checks_list': [
      ('Low Test Class Ratio (<30% of production components)', 'high'),
      ('Classes/Triggers with Zero Test Coverage', 'critical'),
      ('Components Below 75% Test Coverage', 'high'),
      ('Triggers Without a Dedicated Test Class (by naming convention)', 'medium'),
    ],
  },
  {
    'name': 'Org Limits',
    'checks': 5,
    'group': 'Performance & Limits',
    'note': 'The tool queries the Salesforce Limits REST API which returns every org limit. '
            'The 5 check types below represent the distinct finding patterns. Individual limits '
            'monitored include: Daily API Requests, Daily Async Apex Executions, Daily Bulk API Batches, '
            'Daily Workflow Emails, Data Storage (MB), File Storage (MB), Daily Published Platform Events, '
            'Hourly Published Platform Events, Hourly Time-Based Workflow, Permission Sets, Single Email, '
            'Mass Email, Streaming API Concurrent Clients, Monthly Platform Events Entitlement, and all '
            'other limits returned by the API.',
    'checks_list': [
      ('Org Limit ≥90% Consumed — Any limit returned by Salesforce Limits API', 'critical'),
      ('Org Limit 75–89% Consumed — Any limit returned by Salesforce Limits API', 'high'),
      ('Org Limit 50–74% Consumed — Any limit returned by Salesforce Limits API', 'medium'),
      ('Active Apex Classes Approaching Org Limit (>4,500 = High; >4,000 = Medium)', 'high'),
      ('Custom Object Count Approaching Org Limit (>800 = High; >600 = Medium)', 'high'),
    ],
  },
  {
    'name': 'Duplicate & Matching Rules',
    'checks': 4,
    'group': 'Configuration & Architecture',
    'checks_list': [
      ('No Duplicate Rules Configured (includes Salesforce standard rules)', 'high'),
      ('Inactive Duplicate Rules', 'medium'),
      ('No Matching Rules Configured (includes Salesforce standard rules)', 'high'),
      ('Duplicate Rules Without Descriptions', 'low'),
    ],
  },
  {
    'name': 'Reports & Dashboards',
    'checks': 3,
    'group': 'Governance & Hygiene',
    'checks_list': [
      ('Reports Not Run in 6+ Months (>50 stale reports)', 'medium'),
      ('Dashboards Not Viewed in 6+ Months (>20 stale dashboards)', 'medium'),
      ('Total Reports in Org (>2,000)', 'medium'),
    ],
  },
  {
    'name': 'Email Templates',
    'checks': 3,
    'group': 'Governance & Hygiene',
    'checks_list': [
      ('Classic (Non-Lightning) Email Templates Present', 'medium'),
      ('Email Templates Not Modified in 2+ Years', 'low'),
      ('No Email Templates Found in Org', 'low'),
    ],
  },
  {
    'name': 'Platform Events & CDC',
    'checks': 3,
    'group': 'Performance & Limits',
    'checks_list': [
      ('Platform Event Channels with No Active Subscribers', 'high'),
      ('Change Data Capture Entities Enabled (>20)', 'medium'),
      ('No Custom Platform Events, CDC, or Managed-Package Events Configured', 'low'),
    ],
  },
  {
    'name': 'Managed Packages',
    'checks': 3,
    'group': 'CRM & Service',
    'checks_list': [
      ('Managed Packages Installed (>20)', 'medium'),
      ('Beta Managed Packages Installed in Production', 'high'),
      ('Managed Packages — Review Version Currency', 'low'),
    ],
  },
  {
    'name': 'Custom Metadata & Settings',
    'checks': 3,
    'group': 'Configuration & Architecture',
    'checks_list': [
      ('Custom Settings in Use — Legacy Configuration Pattern (Should Use Custom Metadata Types)', 'medium'),
      ('Custom Settings Without Descriptions', 'low'),
      ('No Custom Metadata Types Found — All Configuration Uses Custom Settings', 'medium'),
    ],
  },
  {
    'name': 'Record Types & Page Layouts',
    'checks': 4,
    'group': 'Configuration & Architecture',
    'checks_list': [
      ('Inactive Record Types Present', 'medium'),
      ('Custom Record Types Across Org (>100)', 'medium'),
      ('Record Types Without Descriptions (>10)', 'low'),
      ('Page Layouts Configured (>100)', 'medium'),
    ],
  },
  {
    'name': 'Einstein & AI Usage',
    'checks': 9,
    'group': 'CRM & Service',
    'checks_list': [
      ('Einstein Generative AI / Agentforce Not Enabled (multi-signal detection)', 'low'),
      ('Einstein Prediction Builder Enabled but No Prompt Templates Configured', 'medium'),
      ('Inactive Bot/Agent Definitions', 'medium'),
      ('Einstein/Agentforce Enabled but No Implementation Found (no bots, no prompts)', 'low'),
      ('Einstein AI Applications Exist but None Are Active — License Unused', 'medium'),
      ('Einstein Case Classification Active but Insufficient Training Data (<1,000 closed cases)', 'high'),
      ('Agentforce Bots Exist but No Agent Topics Configured', 'high'),
      ('Agentforce Topics Configured but No Agent Actions Defined', 'high'),
      ('Agentforce Active but Data Cloud Not Connected', 'medium'),
    ],
  },
  {
    'name': 'Experience Cloud',
    'checks': 15,
    'group': 'Security & Access',
    'checks_list': [
      ('Experience Sites Using Legacy Aura Template (aloha, kokua, nto, oob, partner_central)', 'high'),
      ('Live Sites with Unknown/Undetectable Template Type', 'low'),
      ('Inactive/Draft Experience Cloud Sites', 'low'),
      ('Sites with Self-Registration Enabled', 'medium'),
      ('Active Sites with Guest User Access', 'medium'),
      ('Active Sites Without a Custom Domain (still on *.force.com)', 'low'),
      ('Active Sites — Governance Review Recommended (>5 active sites)', 'medium'),
      ('Live Sites Without CDN Enabled', 'low'),
      ('Custom Domains Without HTTPS Enforced', 'high'),
      ("WCAG 2.2 Accessibility Release Updates Not Enabled — Enforced Summer '26", 'medium'),
      ('Active Sites with Clickjack Protection Disabled (AllowAll)', 'critical'),
      ('Live Sites with Browser XSS Protection Disabled', 'medium'),
      ('Live Sites with Content Sniffing Protection Disabled', 'medium'),
      ('Aura Sites with Guest Page Caching Disabled (GuestCacheMaxAge = 0)', 'high'),
      ('Sites with More Than 30 Experience Builder Pages', 'medium'),
    ],
  },
  {
    'name': 'Connected App Security',
    'checks': 12,
    'group': 'Security & Access',
    'checks_list': [
      ('Connected Apps Without Session Timeout', 'high'),
      ('Connected Apps Without Descriptions', 'medium'),
      ('Connected Apps with 20+ Active OAuth Tokens', 'medium'),
      ('Stale OAuth Tokens Not Used in 90+ Days', 'high'),
      ('Duplicate Connected App Names Detected', 'medium'),
      ('Total Connected Apps in Org (>30) — Governance Flag', 'low'),
      ('Active Outbound Messages — Session ID Auth Retired February 2026', 'high'),
      ('Certificates Exceed 200-Day Maximum Lifespan (March 2026 cap)', 'medium'),
      ('CTI / Telephony Connected Apps Without Session Timeout', 'medium'),
      ("Traditional Connected Apps Without External Client App Equivalents (Spring '26 Standard)", 'medium'),
      ('Active OAuth Tokens Belonging to Deactivated Users', 'high'),
      ('Connected Apps Bypassing IP Login Restrictions (IpRelaxation = RelaxedForThisApp)', 'medium'),
    ],
  },
  {
    'name': 'LWC & Components',
    'checks': 39,
    'group': 'Code & Development',
    'checks_list': [
      ('LWC Bundles Without Descriptions', 'low'),
      ('LWC Bundles on Outdated API Versions (< v57)', 'medium'),
      ("LWC Bundles on Retired API Versions (≤ v30) — Broken Since Summer '25", 'critical'),
      ('Aura Components vs LWC — >40% Aura (Migration Debt)', 'medium'),
      ('Aura Components with No LWC Migration Started (zero LWC bundles)', 'high'),
      ('Aura Components with Custom RENDERER Definitions', 'medium'),
      ('Aura Application/Component Events Defined (>5)', 'low'),
      ('LWC Bundles Without Jest Test Files', 'medium'),
      ('LWC/Aura Components Not Modified in 2+ Years', 'low'),
      ('Managed Package LWC Components Modified Locally (will be overwritten on upgrade)', 'medium'),
      ('LWC Components Contain debugger Statements — no-debugger', 'critical'),
      ('LWC Components Use .innerHTML Assignment — XSS Risk — no-inner-html', 'high'),
      ('LWC Components Query the Document Directly — Shadow DOM Bypass — no-document-query', 'high'),
      ('LWC Components Add Event Listeners Without Removing Them — Memory Leak — no-leaky-event-listeners', 'medium'),
      ('LWC Components Use Async Timer Operations (setTimeout/setInterval) — no-async-operation', 'medium'),
      ('LWC Components Use async/await — no-async-await', 'medium'),
      ('LWC Components Reference Browser Globals (window/navigator/location) — SSR Incompatible', 'medium'),
      ('LWC Components Use for...of Loops — no-for-of', 'low'),
      ('LWC Components Use Rest Parameters (...args) — no-rest-parameter', 'low'),
      ('LWC Components Reference process.env.NODE_ENV — no-node-env-in-ssr', 'low'),
      ('LWC Components Have Duplicate Import Statements — no-duplicate-imports', 'low'),
      ('LWC Components Use eval() — Critical Security Vulnerability', 'critical'),
      ('LWC Components Contain console Statements in Production', 'medium'),
      ('LWC Components Use Deprecated @track Decorator (Redundant Since Spring \'20)', 'low'),
      ('LWC Components Use JSON.parse(JSON.stringify()) Deep Clone Anti-Pattern', 'medium'),
      ('LWC Components Mutate Inline Styles via JavaScript (.style.)', 'low'),
      ('LWC Component JS Files Exceed 500 Lines', 'medium'),
      ('LWC Components Use Deprecated if:true / if:false Directives (Removed v60)', 'medium'),
      ('LWC Components Use for:each Without a key Attribute — Full Re-Render on Every Change', 'high'),
      ('LWC Templates Use Inline style= Attributes', 'low'),
      ('LWC Components Fire CustomEvent with bubbles:true AND composed:true — Shadow DOM Breach', 'high'),
      ('LWC Components Override SLDS Classes with Hardcoded Hex/RGB Colors', 'medium'),
      ('Lightning Record Pages Not Modified in 2+ Years', 'low'),
      ('High Total Lightning Page Count (>50) — Governance Flag', 'low'),
      ('Objects with 4+ Lightning Record Pages', 'low'),
      ('Visualforce Pages Present in Org — Legacy UI Technology', 'medium'),
      ('Visualforce Pages on Old API Versions (< v50)', 'high'),
      ('Visualforce Pages Without Descriptions', 'low'),
      ('Visualforce Pages Enabled for Salesforce Mobile (IsAvailableInTouch) — Poor UX', 'medium'),
    ],
  },
  {
    'name': 'OmniStudio',
    'checks': 26,
    'group': 'Code & Development',
    'note': 'Automatically detects native OmniStudio (OmniProcess) or managed package Vlocity '
            '(vlocity_cmt__, vlocity_ins__, vlocity_ps__). Returns 100% Healthy if OmniStudio is not installed.',
    'checks_list': [
      ('Active OmniScripts in Test Mode — Exposes Debug Info to End Users', 'critical'),
      ('Active Integration Procedures in Test Mode', 'critical'),
      ('OmniScript Types with Multiple Active Versions (Version Sprawl)', 'high'),
      ('Active OmniScripts Without LWC Compilation Enabled (Aura Runtime Still Active)', 'medium'),
      ('Active OmniScripts Using Deprecated Remote Action Elements', 'high'),
      ('Active Integration Procedures With No Error-Handling Elements (SetErrors/Throw)', 'high'),
      ('Legacy Knowledge Article Types Still in Schema (Pre-Spring \'20 Migration)', 'high'),
      ('Inactive OmniScripts', 'medium'),
      ('Inactive Integration Procedures', 'medium'),
      ('Inactive DataRaptors / Data Transforms', 'medium'),
      ('Extract DataRaptors Without Turbo Extract Enabled', 'medium'),
      ('High DataRaptor / Data Transform Volume (>100)', 'medium'),
      ('Outdated OmniStudio Managed Package Version (Managed Package Orgs)', 'medium'),
      ('OmniScript Naming Convention Violations (Spaces in Type/SubType)', 'medium'),
      ('Over-Reliance on Standard Extract vs Turbo Extract DataTransforms', 'medium'),
      ('OmniScripts Without Descriptions', 'low'),
      ('Integration Procedures Without Descriptions', 'low'),
      ('DataRaptors / Data Transforms Without Descriptions', 'low'),
      ('FlexCards Without Descriptions', 'low'),
      ('OmniScripts Not Modified in 2+ Years', 'low'),
      ('Integration Procedures Not Modified in 2+ Years', 'low'),
      ('DataRaptors / Data Transforms Not Modified in 2+ Years', 'low'),
      ('Inactive FlexCards', 'low'),
      ('FlexCards Not Modified in 2+ Years', 'low'),
      ('Active Integration Procedures With No Active OmniScripts Referencing Them', 'low'),
      ('DataTransform Type Distribution — Standard Extract vs Turbo Extract Breakdown', 'low'),
    ],
  },
  {
    'name': 'Performance',
    'checks': 22,
    'group': 'Performance & Limits',
    'checks_list': [
      ('Objects with Multiple Active Triggers (>1 trigger per object)', 'high'),
      ('Async Apex Failures in the Last 30 Days', 'high'),
      ('Batch Apex Jobs Running Concurrently (Processing + Holding)', 'high'),
      ('Active Debug Trace Flags', 'high'),
      ('Async Apex Jobs Queued (BatchApex + Queueable + Future > threshold)', 'high'),
      ('Future/Queueable Jobs Pending', 'medium'),
      ('Async Jobs Stuck in Processing for 24+ Hours', 'high'),
      ('Active Flows with DML Elements Inside Loops', 'high'),
      ('Apex Classes Exceed 5,000 Lines', 'high'),
      ('Total Active Flows Exceeds 300', 'medium'),
      ('Scheduled Apex Jobs Waiting (>20)', 'medium'),
      ('Active Record-Triggered Flows on Same Object (>3 flows per object)', 'medium'),
      ('Active Scheduled Flows (>10)', 'medium'),
      ('Custom Objects with 300+ Fields (Wide Objects)', 'medium'),
      ('Large Static Resources Over 500 KB', 'medium'),
      ('Apex Classes Over 1,000 Lines', 'medium'),
      ('Aura Component Bundles Still Deployed (>50)', 'low'),
      ('Objects with 6+ Lightning Pages', 'low'),
      ('Platform Cache Not Configured', 'low'),
      ('Event Monitoring Not Producing Logs (Last 7 Days)', 'low'),
      ('Obsolete Flow Versions in Org (>200)', 'low'),
      ('Aura Components Still Deployed — Migration Debt (>50)', 'low'),
    ],
  },
  {
    'name': 'Notes & Attachments',
    'checks': 12,
    'group': 'Governance & Hygiene',
    'checks_list': [
      ('Legacy Note Records Found (classic Notes object)', 'high'),
      ('Legacy Attachment Records Found (classic Attachments object)', 'high'),
      ('Files Permanently Shared Externally (ContentDistribution with No Expiry)', 'high'),
      ('Files Shared Externally via Content Delivery (any expiry)', 'high'),
      ('Enhanced Notes Not Enabled', 'medium'),
      ('Orphaned ContentDocument Records (no ContentDocumentLink)', 'medium'),
      ('Files Larger Than 25 MB', 'medium'),
      ('Objects with 10,000+ File Attachments', 'medium'),
      ('Files With No Title', 'low'),
      ('No Salesforce Content Libraries Configured', 'low'),
      ('Files Not Viewed in 2+ Years', 'low'),
      ('File Distribution Across Objects (Top 20 by attachment count)', 'low'),
    ],
  },
]

# ── Helpers ──────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def sev_hex_bg(sev):
    return {
        'critical': 'FDECEA',
        'high':     'FEF3E2',
        'medium':   'FEFCE7',
        'low':      'EAFAF1',
    }[sev]

def sev_hex_text(sev):
    return {
        'critical': 'C0392B',
        'high':     'D35400',
        'medium':   'B8860B',
        'low':      '27AE60',
    }[sev]

# ── Build document ────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# ── Title page block ────────────────────────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run('SF Tech Debt Assessor')
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(0x03, 0x2D, 0x60)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Complete Check Reference  —  All 22 Categories  —  June 17, 2026')
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
r.font.italic = True

by = doc.add_paragraph()
by.alignment = WD_ALIGN_PARAGRAPH.CENTER
rb = by.add_run('By Steven Bilgram, Success Architect  |  sf-tech-debt-assessor.onrender.com')
rb.font.size = Pt(10)
rb.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

doc.add_paragraph()

# Summary note
note = doc.add_paragraph()
rn = note.add_run(
    'This document lists every individual check across all 22 assessment categories. '
    'Checks are marked Critical, High, Medium, or Low. The tool runs all checks on every '
    'assessment and only surfaces findings where the check condition is met.'
)
rn.font.size = Pt(10)
rn.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
note.paragraph_format.space_after = Pt(12)

doc.add_paragraph()

# ── Category totals summary table ────────────────────────────────────
sum_heading = doc.add_paragraph()
sh = sum_heading.add_run('Category Summary')
sh.font.bold = True
sh.font.size = Pt(13)
sh.font.color.rgb = RGBColor(0x03, 0x2D, 0x60)

tbl = doc.add_table(rows=1, cols=4)
tbl.style = 'Table Grid'
hdr = tbl.rows[0].cells
for cell, txt in zip(hdr, ['Category', 'Group', 'Checks', 'Severities']):
    cell.text = txt
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    set_cell_bg(cell, '032D60')
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

total_checks = sum(c['checks'] for c in CATEGORIES)

for cat in CATEGORIES:
    sev_counts = {}
    for _, s in cat['checks_list']:
        sev_counts[s] = sev_counts.get(s, 0) + 1
    parts = []
    for sev in ['critical', 'high', 'medium', 'low']:
        if sev in sev_counts:
            parts.append(f"{sev_counts[sev]} {sev.capitalize()}")
    row = tbl.add_row().cells
    row[0].text = cat['name']
    row[0].paragraphs[0].runs[0].font.size = Pt(9)
    row[1].text = cat['group']
    row[1].paragraphs[0].runs[0].font.size = Pt(9)
    row[2].text = str(cat['checks'])
    row[2].paragraphs[0].runs[0].font.size = Pt(9)
    row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row[3].text = '  ·  '.join(parts)
    row[3].paragraphs[0].runs[0].font.size = Pt(9)

# Total row
tot = tbl.add_row().cells
tot[0].text = 'TOTAL'
tot[0].paragraphs[0].runs[0].font.bold = True
tot[0].paragraphs[0].runs[0].font.size = Pt(9)
tot[2].text = str(total_checks)
tot[2].paragraphs[0].runs[0].font.bold = True
tot[2].paragraphs[0].runs[0].font.size = Pt(9)
tot[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
for cell in tot:
    set_cell_bg(cell, 'EAF4FB')

doc.add_page_break()

# ── Per-category detail ───────────────────────────────────────────────
for cat in CATEGORIES:
    # Category heading
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(4)
    h.paragraph_format.space_after  = Pt(2)
    hr = h.add_run(f"{cat['name']}  ·  {cat['checks']} checks  ·  {cat['group']}")
    hr.font.bold = True
    hr.font.size = Pt(14)
    hr.font.color.rgb = RGBColor(0x03, 0x2D, 0x60)

    # Optional note
    if 'note' in cat:
        np_ = doc.add_paragraph()
        nr = np_.add_run(f"ℹ  {cat['note']}")
        nr.font.size = Pt(9)
        nr.font.italic = True
        nr.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
        np_.paragraph_format.space_after = Pt(4)

    # Checks table
    ct = doc.add_table(rows=1, cols=3)
    ct.style = 'Table Grid'

    # Header row
    ch = ct.rows[0].cells
    for cell, txt in zip(ch, ['#', 'Check', 'Severity']):
        cell.text = txt
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_bg(cell, '1A5276')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Set column widths
    for i, width in enumerate([Cm(1.0), Cm(13.5), Cm(2.5)]):
        for cell in ct.columns[i].cells:
            cell.width = width

    for idx, (check_name, sev) in enumerate(cat['checks_list'], 1):
        row = ct.add_row().cells
        row[0].text = str(idx)
        row[0].paragraphs[0].runs[0].font.size = Pt(9)
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        row[1].text = check_name
        row[1].paragraphs[0].runs[0].font.size = Pt(9.5)

        row[2].text = SEV_LABEL[sev]
        row[2].paragraphs[0].runs[0].font.size = Pt(8.5)
        row[2].paragraphs[0].runs[0].font.bold = True
        row[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(
            int(sev_hex_text(sev)[0:2], 16),
            int(sev_hex_text(sev)[2:4], 16),
            int(sev_hex_text(sev)[4:6], 16),
        )
        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(row[2], sev_hex_bg(sev))

        if idx % 2 == 0:
            set_cell_bg(row[0], 'F4F6F7')
            set_cell_bg(row[1], 'F4F6F7')

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    doc.add_page_break()

# ── Save ─────────────────────────────────────────────────────────────
out = os.path.expanduser(
    '~/Desktop/SF_Tech_Debt_Assessor_All_Checks_2026-06-17.docx'
)
doc.save(out)
print(f"Saved: {out}")
print(f"Total checks documented: {total_checks}")
