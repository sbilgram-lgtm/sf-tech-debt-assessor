from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

README = "/Users/sbilgram/sf-tech-debt-assessor/README.md"
OUTPUT = "/Users/sbilgram/Desktop/SF_Tech_Debt_Assessor_README.docx"

doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

BLUE  = RGBColor(0x24, 0x78, 0xDB)
DARK  = RGBColor(0x2C, 0x3E, 0x50)
GREY  = RGBColor(0x7F, 0x8C, 0x8D)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TH_BG = RGBColor(0x2C, 0x3E, 0x50)
ROW_ALT = RGBColor(0xF8, 0xF9, 0xFA)
CODE_GREEN = RGBColor(0x27, 0xAE, 0x60)
CODE_RED   = RGBColor(0xC0, 0x39, 0x2B)

def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)

def add_inline(p, text):
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')
    for part in pattern.split(text):
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(10)
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
            run.font.size = Pt(10)
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = CODE_RED
        else:
            run = p.add_run(part)
            run.font.size = Pt(10)

def add_heading(text, level):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = BLUE if level == 1 else DARK
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)

def add_code_line(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text if text else ' ')
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = CODE_GREEN
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F4F4F4')
    pPr.append(shd)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr_cells[i], TH_BG)
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, text in enumerate(row_data):
            if r_idx % 2 == 1:
                set_cell_bg(cells[c_idx], ROW_ALT)
            p = cells[c_idx].paragraphs[0]
            add_inline(p, text)
            for run in p.runs:
                run.font.size = Pt(9)
    doc.add_paragraph()

# ── Parse ────────────────────────────────────────────────────────
with open(README) as f:
    lines = [l.rstrip('\n') for l in f.readlines()]

in_code  = False
in_table = False
table_headers = []
table_rows    = []

for line in lines:
    # Code fence
    if line.startswith('```'):
        in_code = not in_code
        continue
    if in_code:
        add_code_line(line)
        continue

    # Table
    if line.startswith('|'):
        cols = [c.strip() for c in line.strip('|').split('|')]
        if all(re.match(r'^[-: ]+$', c) for c in cols):
            continue
        if not in_table:
            in_table = True
            table_headers = cols
            table_rows = []
        else:
            table_rows.append(cols)
        continue
    else:
        if in_table:
            add_table(table_headers, table_rows)
            in_table = False

    # Headings
    if line.startswith('#### '): add_heading(line[5:], 4); continue
    if line.startswith('### '):  add_heading(line[4:], 3); continue
    if line.startswith('## '):   add_heading(line[3:], 2); continue
    if line.startswith('# '):    add_heading(line[2:], 1); continue

    # Horizontal rule
    if re.match(r'^---+$', line):
        p = doc.add_paragraph('─' * 72)
        p.runs[0].font.color.rgb = GREY
        p.runs[0].font.size = Pt(8)
        continue

    # Blockquote
    if line.startswith('> '):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(line[2:])
        run.italic = True
        run.font.color.rgb = GREY
        run.font.size = Pt(9)
        continue

    # Bullet
    if re.match(r'^[-*] ', line):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        add_inline(p, line[2:])
        continue

    # Numbered list
    m = re.match(r'^(\d+)\. ', line)
    if m:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(2)
        add_inline(p, line[len(m.group(0)):])
        continue

    # Empty line
    if not line.strip():
        continue

    # Normal paragraph
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_inline(p, line)

if in_table:
    add_table(table_headers, table_rows)

doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
